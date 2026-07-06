#!/usr/bin/env python3
"""
Extract KT (комплексное тестирование) exam questions from images / PDFs / DOCX
into structured JSON using the Claude API (Message Batches, 50% cheaper).

Usage:
    python extract_questions.py ./materials                 # process everything
    python extract_questions.py ./materials --limit 5       # trial run on 5 items
    python extract_questions.py ./materials --model claude-haiku-4-5

Requires ANTHROPIC_API_KEY in the environment (or an `ant auth login` profile).
Re-running is safe: already-processed items are skipped, unfinished batches
are resumed.
"""

import argparse
import base64
import hashlib
import io
import json
import sys
import time
import zipfile
from pathlib import Path

import anthropic

IMAGE_EXTS = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}

MAX_IMAGE_BYTES = 4_000_000        # downscale images above this (API limit ~5MB)
MAX_IMAGE_EDGE = 2576              # max useful resolution on Opus 4.8
MAX_PDF_BYTES = 22_000_000         # base64 inflation must stay under 32MB request cap
MAX_DOCX_CHARS = 18_000            # chunk longer DOCX text across several requests
MAX_BATCH_BYTES = 150_000_000      # stay under the 256MB per-batch cap
MAX_BATCH_REQUESTS = 300
POLL_SECONDS = 60

SYSTEM_PROMPT = """\
You extract exam questions from study materials for КТ (комплексное тестирование).
The materials are in Russian and/or Kazakh — preserve the original language exactly,
do not translate.

Extract EVERY question visible in the provided material. For each question:
- "question": the full question text
- "options": all answer options in their original order, WITHOUT the letter/number
  prefixes (strip "A)", "В)", "1.", etc.)
- "correct_answer_indices": the 0-based indices of ALL options marked as correct —
  checkmark, bold, underline, highlight, "+", asterisk, an answer key, a different
  color. Some questions have several correct answers: include every marked index.
  If nothing marks any correct answer, use an empty list. Never guess.
- "topic": the subject/topic if stated in the material, otherwise null
- "needs_review": true if the text is hard to read, options look cut off or
  incomplete, or the correct-answer marking is ambiguous

If the material contains no questions, return an empty "questions" list.\
"""

USER_PROMPT = "Extract all exam questions from this material as JSON."

OUTPUT_SCHEMA = {
    "type": "object",
    "properties": {
        "questions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "question": {"type": "string"},
                    "options": {"type": "array", "items": {"type": "string"}},
                    "correct_answer_indices": {
                        "type": "array",
                        "items": {"type": "integer"},
                    },
                    "topic": {"anyOf": [{"type": "string"}, {"type": "null"}]},
                    "needs_review": {"type": "boolean"},
                },
                "required": [
                    "question",
                    "options",
                    "correct_answer_indices",
                    "topic",
                    "needs_review",
                ],
                "additionalProperties": False,
            },
        }
    },
    "required": ["questions"],
    "additionalProperties": False,
}


# ---------------------------------------------------------------- input units

def image_block(data: bytes, media_type: str) -> dict:
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": base64.standard_b64encode(data).decode(),
        },
    }


def prepare_image(data: bytes, media_type: str, label: str) -> dict | None:
    """Downscale/recompress oversized images so the API accepts them."""
    needs_resize = len(data) > MAX_IMAGE_BYTES
    if not needs_resize:
        return image_block(data, media_type)
    try:
        from PIL import Image
    except ImportError:
        print(f"  ! {label}: {len(data) // 1024}KB is too large and Pillow is not "
              f"installed — skipping. Run: pip install pillow")
        return None
    img = Image.open(io.BytesIO(data))
    if max(img.size) > MAX_IMAGE_EDGE:
        img.thumbnail((MAX_IMAGE_EDGE, MAX_IMAGE_EDGE))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return image_block(buf.getvalue(), "image/jpeg")


def docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def docx_images(path: Path):
    """Yield (name, bytes, media_type) for images embedded in a .docx."""
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                ext = Path(name).suffix.lower()
                if ext in IMAGE_EXTS:
                    yield Path(name).name, zf.read(name), IMAGE_EXTS[ext]


def unit_id(source: str) -> str:
    """Stable ID derived from the source label, so adding/removing files in the
    input folder never misaligns the resume state."""
    return "u" + hashlib.sha1(source.encode("utf-8")).hexdigest()[:16]


def chunk_text(text: str, size: int = MAX_DOCX_CHARS, overlap: int = 1200) -> list[str]:
    """Split long text at line boundaries; overlap tails so a question cut at
    a boundary still appears whole in the next chunk (dedup drops the copy)."""
    if len(text) <= size:
        return [text]
    chunks, current = [], ""
    for line in text.split("\n"):
        if current and len(current) + len(line) > size:
            chunks.append(current)
            current = current[-overlap:]
        current = f"{current}\n{line}" if current else line
    if current:
        chunks.append(current)
    return chunks


def collect_units(input_dir: Path) -> list[dict]:
    """One unit = one API request: {id, source, content_blocks}."""
    units = []

    def add(source: str, blocks: list) -> None:
        units.append({"id": unit_id(source), "source": source, "blocks": blocks})

    files = sorted(p for p in input_dir.rglob("*") if p.is_file())
    for path in files:
        rel = str(path.relative_to(input_dir))
        ext = path.suffix.lower()

        if ext in IMAGE_EXTS:
            block = prepare_image(path.read_bytes(), IMAGE_EXTS[ext], rel)
            if block:
                add(rel, [block, {"type": "text", "text": USER_PROMPT}])

        elif ext == ".pdf":
            data = path.read_bytes()
            if len(data) > MAX_PDF_BYTES:
                print(f"  ! {rel}: PDF over {MAX_PDF_BYTES // 1_000_000}MB — "
                      f"split it into smaller files and re-run. Skipping.")
                continue
            doc = {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": base64.standard_b64encode(data).decode(),
                },
            }
            add(rel, [doc, {"type": "text", "text": USER_PROMPT}])

        elif ext == ".docx":
            text = docx_text(path)
            if text.strip():
                chunks = chunk_text(text)
                for ci, chunk in enumerate(chunks, start=1):
                    label = (f"{rel} (text)" if len(chunks) == 1
                             else f"{rel} (text {ci}/{len(chunks)})")
                    add(label, [{
                        "type": "text",
                        "text": f"{USER_PROMPT}\n\n<document>\n{chunk}\n</document>",
                    }])
            for name, data, media in docx_images(path):
                block = prepare_image(data, media, f"{rel}/{name}")
                if block:
                    add(f"{rel} ({name})",
                        [block, {"type": "text", "text": USER_PROMPT}])
    return units


# ------------------------------------------------------------------- batching

def unit_to_request(unit: dict, model: str) -> dict:
    return {
        "custom_id": unit["id"],
        "params": {
            "model": model,
            "max_tokens": 32000,
            "system": SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": unit["blocks"]}],
            "output_config": {
                "format": {"type": "json_schema", "schema": OUTPUT_SCHEMA}
            },
        },
    }


def request_size(req: dict) -> int:
    return sum(
        len(b["source"]["data"]) if "source" in b else len(b.get("text", ""))
        for b in req["params"]["messages"][0]["content"]
    )


def chunk_requests(requests: list[dict]) -> list[list[dict]]:
    chunks, current, size = [], [], 0
    for req in requests:
        rs = request_size(req)
        if current and (size + rs > MAX_BATCH_BYTES or len(current) >= MAX_BATCH_REQUESTS):
            chunks.append(current)
            current, size = [], 0
        current.append(req)
        size += rs
    if current:
        chunks.append(current)
    return chunks


# ---------------------------------------------------------------- state & I/O

def load_state(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {"sources": {}, "done": [], "batches": {}}


def save_state(path: Path, state: dict) -> None:
    path.write_text(json.dumps(state, ensure_ascii=False, indent=2),
                    encoding="utf-8")


def collect_batch_results(client, batch_id: str, state: dict,
                          raw_dir: Path, errors: list) -> None:
    for result in client.messages.batches.results(batch_id):
        cid = result.custom_id
        source = state["sources"].get(cid, cid)
        if result.result.type == "succeeded":
            msg = result.result.message
            text = next((b.text for b in msg.content if b.type == "text"), "")
            try:
                payload = json.loads(text)
            except json.JSONDecodeError:
                errors.append({"id": cid, "source": source,
                               "error": "invalid JSON in response"})
                continue
            (raw_dir / f"{cid}.json").write_text(
                json.dumps({"source_file": source,
                            "questions": payload.get("questions", [])},
                           ensure_ascii=False, indent=2),
                encoding="utf-8")
            if cid not in state["done"]:
                state["done"].append(cid)
        else:
            detail = result.result.type
            if detail == "errored":
                detail = f"errored: {result.result.error}"
            errors.append({"id": cid, "source": source, "error": detail})


def merge_results(raw_dir: Path, out_file: Path) -> tuple[int, int, int]:
    """Combine per-unit results into one questions.json."""
    all_questions, qid = [], 1
    for f in sorted(raw_dir.glob("*.json")):
        data = json.loads(f.read_text(encoding="utf-8"))
        for q in data["questions"]:
            if "correct_answer_indices" not in q:  # migrate pre-multi-answer format
                old = q.pop("correct_answer_index", None)
                q["correct_answer_indices"] = [] if old is None else [old]
            q["id"] = qid
            q["source_file"] = data["source_file"]
            all_questions.append(q)
            qid += 1
    out_file.write_text(
        json.dumps(all_questions, ensure_ascii=False, indent=2),
        encoding="utf-8")
    with_answer = sum(1 for q in all_questions
                      if q["correct_answer_indices"])
    review = sum(1 for q in all_questions if q["needs_review"])
    return len(all_questions), with_answer, review


# ------------------------------------------------------------------------ main

def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[1])
    ap.add_argument("input", type=Path, help="folder with images/PDFs/DOCX")
    ap.add_argument("--output", type=Path, default=Path("output"),
                    help="output folder (default: ./output)")
    ap.add_argument("--model", default="claude-haiku-4-5",
                    help="model ID (default: claude-haiku-4-5, the cheapest; "
                         "use claude-opus-4-8 for higher accuracy)")
    ap.add_argument("--limit", type=int, default=None,
                    help="process only the first N units (for a trial run)")
    args = ap.parse_args()

    if not args.input.is_dir():
        print(f"Input folder not found: {args.input}")
        return 1

    args.output.mkdir(parents=True, exist_ok=True)
    raw_dir = args.output / "raw"
    raw_dir.mkdir(exist_ok=True)
    state_file = args.output / "state.json"
    state = load_state(state_file)
    errors: list[dict] = []

    client = anthropic.Anthropic()

    # 1. Resume any batches left unfinished by a previous run
    for batch_id in list(state["batches"]):
        batch = client.messages.batches.retrieve(batch_id)
        if batch.processing_status == "ended":
            print(f"Collecting results of earlier batch {batch_id}...")
            collect_batch_results(client, batch_id, state, raw_dir, errors)
            del state["batches"][batch_id]
        else:
            print(f"Earlier batch {batch_id} still {batch.processing_status} — "
                  f"will wait for it below.")
    save_state(state_file, state)

    # 2. Build work units and drop everything already done or in flight
    print(f"Scanning {args.input} ...")
    units = collect_units(args.input)
    in_flight = {cid for ids in state["batches"].values() for cid in ids}
    for u in units:
        state["sources"][u["id"]] = u["source"]
    pending = [u for u in units
               if u["id"] not in state["done"] and u["id"] not in in_flight]
    if args.limit:
        pending = pending[:args.limit]
    print(f"{len(units)} units total, {len(state['done'])} already done, "
          f"{len(in_flight)} in flight, {len(pending)} to submit.")

    # 3. Submit new batches
    if pending:
        requests = [unit_to_request(u, args.model) for u in pending]
        for chunk in chunk_requests(requests):
            batch = client.messages.batches.create(requests=chunk)
            state["batches"][batch.id] = [r["custom_id"] for r in chunk]
            print(f"Submitted batch {batch.id} ({len(chunk)} requests).")
            save_state(state_file, state)

    # 4. Poll until every batch has ended, collecting results as they finish
    while state["batches"]:
        time.sleep(POLL_SECONDS)
        for batch_id in list(state["batches"]):
            batch = client.messages.batches.retrieve(batch_id)
            counts = batch.request_counts
            print(f"  {batch_id}: {batch.processing_status} "
                  f"(ok={counts.succeeded} err={counts.errored} "
                  f"processing={counts.processing})")
            if batch.processing_status == "ended":
                collect_batch_results(client, batch_id, state, raw_dir, errors)
                del state["batches"][batch_id]
                save_state(state_file, state)

    save_state(state_file, state)

    # 5. Merge everything into one questions.json
    total, with_answer, review = merge_results(raw_dir, args.output / "questions.json")
    if errors:
        (args.output / "errors.json").write_text(
            json.dumps(errors, ensure_ascii=False, indent=2), encoding="utf-8")

    print()
    print(f"Done. {total} questions -> {args.output / 'questions.json'}")
    print(f"  with a marked correct answer: {with_answer}")
    print(f"  flagged needs_review:         {review}")
    if errors:
        print(f"  failed units:                 {len(errors)} "
              f"(see {args.output / 'errors.json'})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
